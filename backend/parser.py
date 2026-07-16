"""
Resume parsing and analysis engine.
Extracts text from PDF/DOCX, scores ATS compatibility,
detects skills, and generates career-path assessments.
"""

import re
import io
from typing import List, Dict, Tuple, Optional

import pypdf
import docx


# ──────────────────────── Skill Catalog ────────────────────────

SKILL_CATALOG: Dict[str, List[str]] = {
    # Frontend
    "React": ["react", "react.js", "reactjs"],
    "Next.js": ["next.js", "nextjs", "next js"],
    "TypeScript": ["typescript", "ts"],
    "JavaScript": ["javascript", "js", "es6", "es2015"],
    "HTML": ["html", "html5"],
    "CSS": ["css", "css3", "scss", "sass", "less"],
    "Redux": ["redux", "redux toolkit", "rtk"],
    "Vue": ["vue", "vue.js", "vuejs"],
    "Angular": ["angular", "angularjs"],
    "Svelte": ["svelte", "sveltekit"],
    "Webpack": ["webpack"],
    "Vite": ["vite"],
    "Tailwind CSS": ["tailwind", "tailwindcss", "tailwind css"],
    "GraphQL": ["graphql", "gql"],
    "Web Performance": ["web vitals", "core web vitals", "lighthouse", "performance optimization", "web performance"],
    "Design Systems": ["design system", "design systems", "component library", "storybook"],
    "A11y (WCAG)": ["accessibility", "a11y", "wcag", "aria", "screen reader"],
    "Testing (Jest)": ["jest", "testing library", "react testing library", "unit test", "unit testing"],
    "Cypress": ["cypress", "e2e testing", "end-to-end"],
    "REST APIs": ["rest", "rest api", "restful", "api"],

    # Backend
    "Node.js": ["node", "node.js", "nodejs", "express", "expressjs"],
    "Python": ["python", "flask", "django", "fastapi"],
    "Java": ["java", "spring", "spring boot", "springboot"],
    "Go": ["golang", "go lang"],
    "Rust": ["rust", "cargo"],
    "C#": ["c#", "csharp", ".net", "dotnet", "asp.net"],
    "Ruby": ["ruby", "rails", "ruby on rails"],
    "PHP": ["php", "laravel", "symfony"],
    "SQL": ["sql", "mysql", "postgresql", "postgres", "sqlite", "mssql"],
    "MongoDB": ["mongodb", "mongo", "mongoose"],
    "Redis": ["redis", "caching", "edge caching"],
    "Kafka": ["kafka", "event streaming", "message queue", "rabbitmq"],
    "Microservices": ["microservice", "microservices", "service mesh"],

    # Platform / Infra / DevOps
    "Docker": ["docker", "container", "containerization"],
    "Kubernetes": ["kubernetes", "k8s", "helm"],
    "AWS": ["aws", "amazon web services", "ec2", "s3", "lambda", "cloudformation"],
    "GCP": ["gcp", "google cloud", "google cloud platform", "bigquery"],
    "Azure": ["azure", "microsoft azure"],
    "Terraform": ["terraform", "iac", "infrastructure as code"],
    "CI/CD": ["ci/cd", "cicd", "github actions", "jenkins", "circleci", "gitlab ci"],
    "Git": ["git", "github", "gitlab", "bitbucket", "version control"],
    "Linux": ["linux", "ubuntu", "debian", "bash", "shell"],
    "Observability": ["datadog", "opentelemetry", "grafana", "prometheus", "observability", "monitoring"],

    # Data / ML
    "Machine Learning": ["machine learning", "ml", "deep learning", "neural network"],
    "TensorFlow": ["tensorflow", "tf"],
    "PyTorch": ["pytorch"],
    "Data Science": ["data science", "data analysis", "pandas", "numpy", "scipy"],
    "NLP": ["nlp", "natural language processing", "text mining"],
    "Computer Vision": ["computer vision", "opencv", "image processing"],

    # Soft / Leadership
    "System Design": ["system design", "architecture", "high-level design", "hld"],
    "Technical Writing": ["technical writing", "documentation", "rfc"],
    "Agile": ["agile", "scrum", "kanban", "sprint"],
    "Leadership": ["tech lead", "team lead", "engineering manager", "mentoring", "mentorship"],
    "Communication": ["stakeholder", "cross-functional", "presentation"],
}

# ──────────────────────── Action Verbs ────────────────────────

STRONG_ACTION_VERBS = {
    "architected", "built", "created", "delivered", "deployed", "designed",
    "developed", "drove", "eliminated", "engineered", "established",
    "executed", "grew", "implemented", "improved", "increased", "integrated",
    "launched", "led", "managed", "migrated", "optimized", "orchestrated",
    "overhauled", "pioneered", "published", "re-architected", "rebuilt",
    "reduced", "refactored", "replaced", "revamped", "scaled", "shipped",
    "simplified", "spearheaded", "streamlined", "transformed", "unified",
}

WEAK_PHRASES = [
    "responsible for", "worked on", "helped with", "assisted in",
    "involved in", "participated in", "was part of", "contributed to",
    "tasked with", "duties included",
]

# ──────────────────────── Role Definitions ────────────────────

ROLE_SKILL_MAP: Dict[str, List[str]] = {
    "Frontend Engineer": [
        "React", "TypeScript", "JavaScript", "HTML", "CSS", "REST APIs",
        "Git", "Redux", "Testing (Jest)", "Webpack", "Next.js", "GraphQL",
        "Web Performance", "Design Systems", "A11y (WCAG)",
    ],
    "Senior Frontend Engineer": [
        "React", "TypeScript", "JavaScript", "HTML", "CSS", "REST APIs",
        "Git", "Redux", "Testing (Jest)", "Webpack", "Next.js", "GraphQL",
        "Web Performance", "Design Systems", "A11y (WCAG)", "System Design",
    ],
    "Backend Engineer": [
        "Node.js", "Python", "SQL", "REST APIs", "Git", "Docker",
        "Redis", "MongoDB", "Kafka", "Microservices", "CI/CD", "Linux",
        "Testing (Jest)", "System Design",
    ],
    "Senior Backend Engineer": [
        "Node.js", "Python", "SQL", "REST APIs", "Git", "Docker",
        "Redis", "MongoDB", "Kafka", "Microservices", "CI/CD", "Linux",
        "Testing (Jest)", "System Design", "Observability", "AWS",
    ],
    "Full Stack Engineer": [
        "React", "TypeScript", "Node.js", "SQL", "REST APIs", "Git",
        "Docker", "CSS", "MongoDB", "Redis", "CI/CD", "Testing (Jest)",
        "GraphQL", "AWS",
    ],
    "Platform / Infrastructure Engineer": [
        "Docker", "Kubernetes", "AWS", "Terraform", "CI/CD", "Git",
        "Linux", "Observability", "Python", "Microservices", "System Design",
    ],
    "Staff Software Engineer (Backend)": [
        "System Design", "Node.js", "Python", "SQL", "REST APIs", "Git",
        "Docker", "Kubernetes", "AWS", "Kafka", "Microservices", "CI/CD",
        "Observability", "Technical Writing", "Leadership",
    ],
    "DevOps Engineer": [
        "Docker", "Kubernetes", "AWS", "Terraform", "CI/CD", "Git",
        "Linux", "Observability", "Python", "Agile",
    ],
    "Data Scientist": [
        "Python", "Machine Learning", "Data Science", "SQL", "TensorFlow",
        "PyTorch", "NLP", "Git",
    ],
    "Engineering Manager": [
        "Leadership", "Agile", "System Design", "Communication",
        "Technical Writing", "Git", "CI/CD",
    ],
    "Solutions Architect": [
        "AWS", "System Design", "Communication", "Technical Writing",
        "Docker", "Kubernetes", "Terraform", "Microservices",
    ],
}

# ──────────────────────── Text Extraction ────────────────────

def analyze_layout_from_pdf(reader: pypdf.PdfReader, text: str) -> Dict:
    """Analyze PDF for tables or multi-column layout using draw operators and text layout."""
    has_tables = False
    table_reasons = []

    # Heuristic 1: Inspect PDF content streams for vector paths / lines
    draw_operators = 0
    for page in reader.pages:
        content = page.get_contents()
        if content:
            try:
                if isinstance(content, list):
                    data = "\n".join(c.get_data().decode('utf-8', errors='ignore') for c in content)
                else:
                    data = content.get_data().decode('utf-8', errors='ignore')
                # Count path markers: l (lineto), re (rectangle), S (stroke), f (fill)
                draw_operators += len(re.findall(r'\b(l|re|S|f)\b', data))
            except Exception:
                pass

    if draw_operators > 40:
        has_tables = True
        table_reasons.append(f"Detected {draw_operators} vector line drawing elements in PDF, indicating border lines, grids, or tables.")

    # Heuristic 2: Large spacing gaps in text (indicative of table cells or multi-columns)
    multi_space_lines = 0
    for line in text.split("\n"):
        if re.search(r'\s{5,}', line.strip()):
            multi_space_lines += 1

    if multi_space_lines > 12:
        has_tables = True
        table_reasons.append(f"Detected {multi_space_lines} lines containing multi-column tabular spacing layout.")

    return {
        "has_tables": has_tables,
        "reasons": table_reasons,
        "file_format": "PDF"
    }


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract all text from a PDF file."""
    reader = pypdf.PdfReader(io.BytesIO(file_bytes))
    pages = []
    for page in reader.pages:
        text = page.extract_text()
        if text:
            pages.append(text)
    return "\n".join(pages)


def extract_text_from_docx_with_layout(file_bytes: bytes) -> Tuple[str, Dict]:
    """Extract text from a DOCX file, including paragraphs and tables, and detect table layout."""
    document = docx.Document(io.BytesIO(file_bytes))
    lines = []
    has_tables = len(document.tables) > 0
    table_reasons = []

    # Extract paragraph text
    for para in document.paragraphs:
        if para.text.strip():
            lines.append(para.text)

    # Extract table text
    if has_tables:
        table_reasons.append(f"Detected {len(document.tables)} document tables natively in DOCX file structure.")
        for table in document.tables:
            for row in table.rows:
                # Merge cell text with a spacing delimiter
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                # Deduplicate repeated text from merged cells
                cleaned_row = []
                for cell in row_text:
                    if not cleaned_row or cleaned_row[-1] != cell:
                        cleaned_row.append(cell)
                if cleaned_row:
                    lines.append(" | ".join(cleaned_row))

    text = "\n".join(lines)
    return text, {
        "has_tables": has_tables,
        "reasons": table_reasons,
        "file_format": "DOCX"
    }


def extract_text_and_layout(file_bytes: bytes, filename: str) -> Tuple[str, Dict]:
    """Route extraction and return both raw text and detected layout facts."""
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        text = extract_text_from_pdf(file_bytes)
        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            layout = analyze_layout_from_pdf(reader, text)
        except Exception:
            layout = {"has_tables": False, "reasons": [], "file_format": "PDF"}
        return text, layout
    elif ext in ("docx", "doc"):
        return extract_text_from_docx_with_layout(file_bytes)
    else:
        text = file_bytes.decode("utf-8", errors="ignore")
        # Check text-spacing for plain text
        multi_space_lines = sum(1 for line in text.split("\n") if re.search(r'\s{5,}', line.strip()))
        has_tables = multi_space_lines > 12
        reasons = [f"Detected {multi_space_lines} lines of tabular spacing."] if has_tables else []
        return text, {"has_tables": has_tables, "reasons": reasons, "file_format": "TXT"}


def extract_text(file_bytes: bytes, filename: str) -> str:
    """Route extraction based on file extension (kept for backward compatibility)."""
    text, _ = extract_text_and_layout(file_bytes, filename)
    return text


# ──────────────────────── Skill Detection ────────────────────

def detect_skills(text: str) -> List[str]:
    """Return canonical skill names found in the resume text."""
    text_lower = text.lower()
    found = []
    for canonical, aliases in SKILL_CATALOG.items():
        for alias in aliases:
            # Use word-boundary-like match for short aliases
            if len(alias) <= 3:
                pattern = r'(?<![a-zA-Z])' + re.escape(alias) + r'(?![a-zA-Z])'
                if re.search(pattern, text_lower):
                    found.append(canonical)
                    break
            else:
                if alias in text_lower:
                    found.append(canonical)
                    break
    return found


# ──────────────────────── Bullet Analysis ────────────────────

def _get_bullets(text: str) -> List[str]:
    """Split text into bullet-like lines (lines starting with -, •, *, or numbered)."""
    lines = text.split("\n")
    bullets = []
    for line in lines:
        stripped = line.strip()
        if stripped and (
            stripped[0] in "-•*▪▸►"
            or re.match(r'^\d+[\.\)]\s', stripped)
            or len(stripped) > 30  # treat longer lines as potential bullets too
        ):
            bullets.append(stripped)
    return bullets


def _count_quantified(bullets: List[str]) -> int:
    """Count bullets containing numbers/percentages/metrics."""
    count = 0
    for b in bullets:
        if re.search(r'\d+[\%xX]|\d{2,}|\$\d|→|↑|↓', b):
            count += 1
    return count


def _count_strong_verbs(bullets: List[str]) -> int:
    """Count bullets starting with strong action verbs."""
    count = 0
    for b in bullets:
        # Strip leading bullet markers
        cleaned = re.sub(r'^[-•*▪▸►\d\.\)\s]+', '', b).strip()
        first_word = cleaned.split()[0].lower().rstrip("ed").rstrip("ing") if cleaned.split() else ""
        # Check exact match on full first word
        first_word_full = cleaned.split()[0].lower() if cleaned.split() else ""
        if first_word_full in STRONG_ACTION_VERBS:
            count += 1
    return count


def _count_weak_phrases(text: str) -> int:
    """Count occurrences of passive/weak phrases."""
    text_lower = text.lower()
    count = 0
    for phrase in WEAK_PHRASES:
        count += text_lower.count(phrase)
    return count


def _check_formatting(text: str) -> int:
    """Score formatting quality (0-100)."""
    score = 70  # baseline
    # Consistent delimiters (· or | or ,)
    if "·" in text or "|" in text:
        score += 10
    # Proper capitalization patterns for common tech
    if re.search(r'Node\.js|TypeScript|JavaScript|MongoDB|GraphQL|Next\.js', text):
        score += 10
    # Messy capitalization
    if re.search(r'nodejs|Nodejs|javascript|typescript|REACT|DOCKER', text):
        score -= 15
    # Multi-column or table layout (negative for ATS)
    if text.count("\t") > 10:
        score -= 10
    # Consistent date formatting
    if re.search(r'(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}', text):
        score += 5
    return max(0, min(100, score))


# ──────────────────────── ATS Scoring ────────────────────────

def compute_ats_score(text: str, skills: List[str], layout_info: Optional[Dict] = None) -> Dict:
    """
    Return composite ATS score and sub-metrics.
    Returns: { score, breakdown: [{l, v, tone}], strengths, improvements }
    """
    bullets = _get_bullets(text)
    total_bullets = max(len(bullets), 1)

    # --- Sub-scores ---
    # Keywords (how many catalog skills detected)
    keyword_ratio = min(len(skills) / 12, 1.0)
    keywords_score = int(50 + keyword_ratio * 50)

    # Impact (quantification)
    quant_count = _count_quantified(bullets)
    quant_ratio = quant_count / total_bullets
    impact_score = int(40 + min(quant_ratio / 0.6, 1.0) * 60)

    # Formatting
    formatting_score = _check_formatting(text)
    if layout_info and layout_info.get("has_tables"):
        formatting_score = max(0, formatting_score - 20)

    # Composite
    composite = int(keywords_score * 0.40 + impact_score * 0.30 + formatting_score * 0.30)

    def _tone(v: int) -> str:
        if v >= 75:
            return "success"
        if v >= 50:
            return "warning"
        return "danger"

    breakdown = [
        {"l": "Keywords", "v": keywords_score, "tone": _tone(keywords_score)},
        {"l": "Formatting", "v": formatting_score, "tone": _tone(formatting_score)},
        {"l": "Impact", "v": impact_score, "tone": _tone(impact_score)},
    ]

    # --- Strengths ---
    strengths: List[str] = []
    strong_verb_count = _count_strong_verbs(bullets)
    if strong_verb_count >= 3:
        strengths.append(f"Strong action verbs in most bullet points ({strong_verb_count} of {total_bullets} bullets use powerful leading verbs).")
    if quant_ratio >= 0.4:
        strengths.append(f"Quantified impact on {int(quant_ratio*100)}% of achievements — recruiters love measurable outcomes.")
    word_count = len(text.split())
    if word_count <= 700:
        strengths.append("Concise format — easily parseable by ATS engines without truncation.")
    if len(skills) >= 8:
        strengths.append(f"Broad skill coverage — {len(skills)} recognized technologies detected across your resume.")
    if formatting_score >= 75:
        if layout_info and layout_info.get("has_tables"):
            strengths.append("Clean formatting aside from table structures.")
        else:
            strengths.append("Clean formatting with consistent capitalization and delimiter usage.")
    if "github" in text.lower() or "linkedin" in text.lower() or "portfolio" in text.lower():
        strengths.append("Includes relevant profile links (LinkedIn/GitHub/Portfolio) for recruiter follow-up.")
    if not strengths:
        strengths.append("Resume successfully parsed and analyzed.")

    # --- Improvements ---
    improvements: List[Dict] = []
    weak_count = _count_weak_phrases(text)

    if layout_info and layout_info.get("has_tables"):
        reasons_str = " ".join(layout_info.get("reasons", []))
        improvements.append({
            "category": "Formatting",
            "color": "warning",
            "before": "Table/grid layout detected: " + (reasons_str[:80] + "..." if len(reasons_str) > 80 else reasons_str),
            "after": "Reformat to a single-column layout without tables or text grids.",
            "why": "Many standard ATS parsers fail to parse multi-column layouts or tables, leading to scrambled text sequences.",
        })

    if quant_ratio < 0.5:
        improvements.append({
            "category": "Quantification",
            "color": "info",
            "before": "Worked on improving app performance.",
            "after": "Reduced p95 API latency by 42% (820ms → 475ms) by introducing Redis edge caching, improving retention by 8%.",
            "why": "ATS ranks bullets with numeric outcomes ~1.7× higher — recruiters filter for measurable business impact.",
        })

    if weak_count >= 2:
        improvements.append({
            "category": "Action Verbs",
            "color": "accent",
            "before": "Responsible for building the checkout system.",
            "after": "Architected the checkout microservice serving 3M+ monthly transactions across 12 markets.",
            "why": "Passive phrases like 'responsible for' rank low. Strong leading verbs match recruiter keyword filters.",
        })

    if formatting_score < 80 and not (layout_info and layout_info.get("has_tables")):
        improvements.append({
            "category": "Formatting",
            "color": "warning",
            "before": "Skills: React, Nodejs, Docker, aws, TypeScript, mongodb",
            "after": "Skills: React · Node.js · TypeScript · Docker · AWS · MongoDB",
            "why": "Normalized capitalization and delimiter formatting improves keyword parsing accuracy in ATS engines like Greenhouse & Lever.",
        })

    if keyword_ratio < 0.8:
        improvements.append({
            "category": "Keywords",
            "color": "success",
            "before": "Built dashboards for internal teams.",
            "after": "Built analytics dashboards (React, D3, TypeScript) surfacing product KPIs to 40+ internal stakeholders.",
            "why": "Adding stack keywords + audience scale directly aligns to filter queries used by technical recruiters.",
        })

    return {
        "score": composite,
        "breakdown": breakdown,
        "strengths": strengths,
        "improvements": improvements,
    }


# ──────────────────────── Role Matching ────────────────────────

def match_role(skills: List[str], target_role: str) -> Dict:
    """
    Compare detected skills against the target role's requirements.
    Returns: { role, score, matched, missing, certs, projects, templateGuidelines }
    """
    # Find closest role definition
    role_key = _find_closest_role(target_role)
    required = ROLE_SKILL_MAP.get(role_key, ROLE_SKILL_MAP.get("Full Stack Engineer", []))

    skill_set = set(skills)
    matched = [s for s in required if s in skill_set]
    missing = [s for s in required if s not in skill_set]

    score = int((len(matched) / max(len(required), 1)) * 100)

    certs = _suggest_certs(target_role, missing)
    projects = _suggest_projects(target_role, missing)
    guidelines = _get_template_guidelines(target_role)

    return {
        "role": target_role,
        "score": score,
        "matched": matched,
        "missing": missing,
        "certs": certs,
        "projects": projects,
        "templateGuidelines": guidelines,
    }


def _find_closest_role(target: str) -> str:
    """Find the closest matching role key in ROLE_SKILL_MAP."""
    target_lower = target.lower()
    best_key = "Full Stack Engineer"
    best_score = 0
    for key in ROLE_SKILL_MAP:
        key_lower = key.lower()
        # Simple word-overlap scoring
        target_words = set(target_lower.split())
        key_words = set(key_lower.split())
        overlap = len(target_words & key_words)
        if overlap > best_score:
            best_score = overlap
            best_key = key
    # Fallback to keyword matching
    if best_score == 0:
        if "frontend" in target_lower or "front-end" in target_lower or "front end" in target_lower:
            best_key = "Frontend Engineer"
        elif "backend" in target_lower or "back-end" in target_lower or "back end" in target_lower:
            best_key = "Backend Engineer"
        elif "devops" in target_lower or "sre" in target_lower:
            best_key = "DevOps Engineer"
        elif "data" in target_lower or "ml" in target_lower or "machine learning" in target_lower:
            best_key = "Data Scientist"
        elif "platform" in target_lower or "infra" in target_lower:
            best_key = "Platform / Infrastructure Engineer"
        elif "manager" in target_lower or "em" in target_lower:
            best_key = "Engineering Manager"
        elif "architect" in target_lower or "solution" in target_lower:
            best_key = "Solutions Architect"
        elif "staff" in target_lower:
            best_key = "Staff Software Engineer (Backend)"
    return best_key


CERT_SUGGESTIONS: Dict[str, List[str]] = {
    "frontend": [
        "Meta Frontend Developer Professional Certificate",
        "Google UX + Web Vitals Specialization",
        "AWS Certified Cloud Practitioner",
    ],
    "backend": [
        "AWS Certified Solutions Architect – Associate",
        "Google Cloud Professional Cloud Developer",
        "MongoDB Certified Developer Associate",
    ],
    "devops": [
        "AWS Certified DevOps Engineer – Professional",
        "Certified Kubernetes Administrator (CKA)",
        "HashiCorp Certified: Terraform Associate",
    ],
    "data": [
        "Google Professional Machine Learning Engineer",
        "AWS Certified Machine Learning – Specialty",
        "TensorFlow Developer Certificate",
    ],
    "default": [
        "AWS Certified Cloud Practitioner",
        "Google Professional Cloud Developer",
        "Meta Back-End Developer Professional Certificate",
    ],
}

PROJECT_SUGGESTIONS: Dict[str, List[str]] = {
    "frontend": [
        "Build an accessible design system with Radix + Storybook.",
        "Ship a Next.js app with ISR and Core Web Vitals ≥ 95.",
        "Migrate a REST client to GraphQL with codegen + persisted queries.",
    ],
    "backend": [
        "Build a real-time notification service with WebSockets + Redis Pub/Sub.",
        "Design a rate-limited API gateway with Express/Fastify + Redis.",
        "Implement an event-sourced order system with Kafka + Postgres.",
    ],
    "devops": [
        "Set up a multi-env Terraform pipeline deploying to AWS ECS.",
        "Build a GitOps workflow with ArgoCD + Kubernetes.",
        "Create a full observability stack (Prometheus + Grafana + alerting).",
    ],
    "data": [
        "Train and deploy an NLP sentiment classifier with FastAPI + HuggingFace.",
        "Build an end-to-end ML pipeline with MLflow tracking.",
        "Create a real-time anomaly detection dashboard with Streamlit.",
    ],
    "default": [
        "Build a production-grade REST API with auth, rate-limiting, and OpenAPI docs.",
        "Create a CI/CD pipeline that deploys to cloud on every merge.",
        "Ship a full-stack app with SSR, auth, and real-time features.",
    ],
}


def _role_category(role: str) -> str:
    """Map a role string to a broad category."""
    r = role.lower()
    if any(k in r for k in ("frontend", "front-end", "front end", "ui")):
        return "frontend"
    if any(k in r for k in ("backend", "back-end", "back end", "server")):
        return "backend"
    if any(k in r for k in ("devops", "sre", "platform", "infra")):
        return "devops"
    if any(k in r for k in ("data", "ml", "machine learning", "ai")):
        return "data"
    return "default"


def _suggest_certs(role: str, missing: List[str]) -> List[str]:
    cat = _role_category(role)
    return CERT_SUGGESTIONS.get(cat, CERT_SUGGESTIONS["default"])


def _suggest_projects(role: str, missing: List[str]) -> List[str]:
    cat = _role_category(role)
    return PROJECT_SUGGESTIONS.get(cat, PROJECT_SUGGESTIONS["default"])


def _get_template_guidelines(role: str) -> List[Dict[str, str]]:
    """Return the ideal resume section template for the given role."""
    return [
        {"s": "Header", "d": "Name · Title · Location · Email · LinkedIn · GitHub · Portfolio URL"},
        {"s": "Summary (3-4 lines)", "d": "Years of experience, primary stack, and 1 measurable achievement."},
        {"s": "Core Skills", "d": "6-10 skills, delimited by · — grouped: Languages, Frameworks, Tooling."},
        {"s": "Experience", "d": "Reverse-chronological. 3-5 bullets per role. Lead with action verb + metric."},
        {"s": "Projects", "d": "2-3 public projects with live links and stack callouts."},
        {"s": "Education & Certifications", "d": "Degree, institution, year. Certifications with expiry if relevant."},
    ]


# ──────────────────────── Career Explorer ────────────────────────

# Archetype mapping based on dominant skill clusters
ARCHETYPE_MAP = {
    "backend_heavy": {
        "title": "Backend Heavy · System Design Architect",
        "description": (
            "You lean toward distributed systems, API design, and infra-adjacent tooling. "
            "Your strengths compound in teams where reliability, throughput, and platform "
            "leverage matter more than pixel-level UI polish."
        ),
        "trigger_skills": {"Node.js", "Python", "SQL", "Redis", "Kafka", "Docker",
                           "Microservices", "System Design", "AWS", "Kubernetes"},
    },
    "frontend_craftsman": {
        "title": "Frontend Craftsman · UI Experience Specialist",
        "description": (
            "You excel at crafting polished user interfaces with attention to accessibility, "
            "performance, and design-system scalability. Teams benefit from your deep understanding "
            "of browser internals and component architecture."
        ),
        "trigger_skills": {"React", "TypeScript", "CSS", "Next.js", "Design Systems",
                           "Web Performance", "A11y (WCAG)", "Testing (Jest)", "Redux"},
    },
    "full_stack_generalist": {
        "title": "Full-Stack Generalist · Versatile Builder",
        "description": (
            "You bridge frontend and backend with equal fluency. Your versatility makes you "
            "the go-to person for greenfield projects and startups that need end-to-end ownership."
        ),
        "trigger_skills": {"React", "Node.js", "SQL", "Docker", "REST APIs", "Git",
                           "TypeScript", "MongoDB", "CSS"},
    },
    "platform_infra": {
        "title": "Platform & Infrastructure · Reliability Architect",
        "description": (
            "You think in systems, not features. Your strengths lie in CI/CD pipelines, "
            "container orchestration, and infrastructure-as-code — the backbone that lets "
            "product teams ship with confidence."
        ),
        "trigger_skills": {"Docker", "Kubernetes", "Terraform", "AWS", "CI/CD", "Linux",
                           "Observability"},
    },
    "data_ml": {
        "title": "Data & ML Engineer · Intelligence Builder",
        "description": (
            "You transform raw data into actionable intelligence. Your expertise spans "
            "model training, feature engineering, and production ML pipelines."
        ),
        "trigger_skills": {"Machine Learning", "Python", "TensorFlow", "PyTorch",
                           "Data Science", "NLP", "SQL"},
    },
}


def _detect_archetype(skills: List[str]) -> Dict:
    """Find the best-matching career archetype for the given skill set."""
    skill_set = set(skills)
    best_key = "full_stack_generalist"
    best_overlap = 0
    for key, info in ARCHETYPE_MAP.items():
        overlap = len(skill_set & info["trigger_skills"])
        if overlap > best_overlap:
            best_overlap = overlap
            best_key = key
    arch = ARCHETYPE_MAP[best_key]
    # Pick top skills that matched as the "archetype skills" tag list
    matched_trigger = list(skill_set & arch["trigger_skills"])
    extra = [s for s in skills if s not in matched_trigger][:max(0, 6 - len(matched_trigger))]
    display_skills = (matched_trigger + extra)[:6]
    return {
        "title": arch["title"],
        "description": arch["description"],
        "skills": display_skills,
    }


# Career paths organized by archetype
CAREER_PATHS = {
    "backend_heavy": {
        "perfectFits": [
            {
                "title": "Senior Backend Engineer",
                "difficulty": "Very Easy",
                "reasons": [
                    "Strong distributed systems and API design background.",
                    "Database and caching experience covers most JD requirements.",
                    "Track record of system design and production ownership.",
                ],
            },
            {
                "title": "Platform / Infrastructure Engineer",
                "difficulty": "Very Easy",
                "reasons": [
                    "Cloud and IaC stack aligns to platform tooling.",
                    "History of building internal developer tooling & templates.",
                    "Comfortable with observability and containerization.",
                ],
            },
            {
                "title": "Staff Software Engineer (Backend)",
                "difficulty": "Easy",
                "reasons": [
                    "Cross-team system design leadership potential.",
                    "Deep API contract & migration strategy experience.",
                    "Demonstrated high-leverage architectural thinking.",
                ],
            },
        ],
        "pivotableRoles": [
            {
                "title": "Solutions Architect",
                "difficulty": "Medium — Requires Upskilling",
                "bridge": [
                    "Sharpen client-facing communication (workshops, discovery).",
                    "Learn 1-2 hyperscaler reference architectures deeply.",
                    "Contribute to at least 2 public technical case studies.",
                ],
            },
            {
                "title": "Engineering Manager (IC → EM)",
                "difficulty": "Medium — Requires Upskilling",
                "bridge": [
                    "Formalize people-management fundamentals (1:1s, growth plans).",
                    "Own a small team as tech lead for 6 months.",
                    "Complete a management essentials cohort (e.g. LeadDev).",
                ],
            },
        ],
        "notAFit": [
            {
                "title": "Embedded Systems Engineer",
                "difficulty": "Extremely Hard",
                "barriers": [
                    "Requires specialized C/C++ and RTOS experience.",
                    "No hardware / firmware background on resume.",
                    "Most roles want EE or CompE degree credentials.",
                ],
            },
            {
                "title": "Quant Researcher (HFT)",
                "difficulty": "Extremely Hard",
                "barriers": [
                    "Requires PhD-level statistics or applied math.",
                    "No demonstrated low-latency C++/kdb+ portfolio.",
                    "Highly gated interview loops (math + probability).",
                ],
            },
        ],
    },
    "frontend_craftsman": {
        "perfectFits": [
            {
                "title": "Senior Frontend Engineer",
                "difficulty": "Very Easy",
                "reasons": [
                    "Deep React/TypeScript expertise matches 90%+ of JDs.",
                    "Design systems and component library experience is a differentiator.",
                    "Strong testing and performance optimization track record.",
                ],
            },
            {
                "title": "UI/UX Engineer",
                "difficulty": "Very Easy",
                "reasons": [
                    "Accessibility expertise is increasingly in demand.",
                    "Strong visual design sensibility in component work.",
                    "Cross-functional experience with design teams.",
                ],
            },
            {
                "title": "Design Systems Lead",
                "difficulty": "Easy",
                "reasons": [
                    "Direct experience building shared component libraries.",
                    "Strong TypeScript + a11y foundation.",
                    "Demonstrated ability to drive adoption across teams.",
                ],
            },
        ],
        "pivotableRoles": [
            {
                "title": "Full-Stack Engineer",
                "difficulty": "Medium — Requires Upskilling",
                "bridge": [
                    "Add backend fundamentals (Node.js/Express or Python/FastAPI).",
                    "Learn database design (SQL + NoSQL basics).",
                    "Build 1-2 end-to-end projects with auth and APIs.",
                ],
            },
            {
                "title": "Product Engineer",
                "difficulty": "Medium — Requires Upskilling",
                "bridge": [
                    "Develop product thinking and user-research skills.",
                    "Ship a feature end-to-end, owning scope and metrics.",
                    "Practice stakeholder communication and roadmap influence.",
                ],
            },
        ],
        "notAFit": [
            {
                "title": "Site Reliability Engineer",
                "difficulty": "Extremely Hard",
                "barriers": [
                    "Requires deep Linux, networking, and infrastructure knowledge.",
                    "No demonstrated on-call / incident response experience.",
                    "Container orchestration and IaC are not in your stack.",
                ],
            },
            {
                "title": "Machine Learning Engineer",
                "difficulty": "Extremely Hard",
                "barriers": [
                    "Requires strong math/statistics foundation.",
                    "No ML framework experience (TensorFlow/PyTorch).",
                    "Typically requires advanced degree or research background.",
                ],
            },
        ],
    },
    # Fallback for other archetypes
    "default": {
        "perfectFits": [
            {
                "title": "Full-Stack Software Engineer",
                "difficulty": "Easy",
                "reasons": [
                    "Broad skill set covers both frontend and backend.",
                    "Versatility is highly valued at startups and growing teams.",
                    "Demonstrated ability to work across the stack.",
                ],
            },
            {
                "title": "Software Engineer II",
                "difficulty": "Very Easy",
                "reasons": [
                    "Solid fundamentals in core languages and frameworks.",
                    "Good testing and version control practices.",
                    "Ready for mid-level ownership with mentorship.",
                ],
            },
            {
                "title": "Backend Engineer",
                "difficulty": "Easy",
                "reasons": [
                    "Strong foundation in server-side technologies.",
                    "Database and API experience aligns with most roles.",
                    "Comfortable with deployment and DevOps basics.",
                ],
            },
        ],
        "pivotableRoles": [
            {
                "title": "DevOps / Platform Engineer",
                "difficulty": "Medium — Requires Upskilling",
                "bridge": [
                    "Deep-dive into Docker, Kubernetes, and Terraform.",
                    "Get certified in a major cloud provider (AWS/GCP/Azure).",
                    "Build and maintain a CI/CD pipeline from scratch.",
                ],
            },
            {
                "title": "Technical Product Manager",
                "difficulty": "Medium — Requires Upskilling",
                "bridge": [
                    "Develop product strategy and user research skills.",
                    "Practice stakeholder management and prioritization.",
                    "Complete a PM certification or bootcamp.",
                ],
            },
        ],
        "notAFit": [
            {
                "title": "Embedded Systems Engineer",
                "difficulty": "Extremely Hard",
                "barriers": [
                    "Requires specialized C/C++ and RTOS experience.",
                    "No hardware / firmware background detected.",
                    "Most roles require EE or CompE degree credentials.",
                ],
            },
            {
                "title": "Quant Researcher (HFT)",
                "difficulty": "Extremely Hard",
                "barriers": [
                    "Requires PhD-level statistics or applied math.",
                    "No demonstrated low-latency C++/kdb+ portfolio.",
                    "Highly gated interview process (math + probability).",
                ],
            },
        ],
    },
}

# Copy default paths for remaining archetypes
for _k in ("full_stack_generalist", "platform_infra", "data_ml"):
    if _k not in CAREER_PATHS:
        CAREER_PATHS[_k] = CAREER_PATHS["default"]


def explore_career(skills: List[str]) -> Dict:
    """Generate full career explorer response."""
    archetype = _detect_archetype(skills)

    # Determine which career path set to use
    skill_set = set(skills)
    best_key = "default"
    best_overlap = 0
    for key, info in ARCHETYPE_MAP.items():
        overlap = len(skill_set & info["trigger_skills"])
        if overlap > best_overlap:
            best_overlap = overlap
            best_key = key

    paths = CAREER_PATHS.get(best_key, CAREER_PATHS["default"])

    return {
        "archetype": archetype,
        "perfectFits": paths["perfectFits"],
        "pivotableRoles": paths["pivotableRoles"],
        "notAFit": paths["notAFit"],
    }


# ──────────────────────── Stats Helpers ────────────────────────

def get_resume_stats(text: str, file_bytes: bytes, filename: str) -> Dict:
    """Return basic stats about the resume."""
    word_count = len(text.split())
    page_count = 1
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else ""
    if ext == "pdf":
        try:
            reader = pypdf.PdfReader(io.BytesIO(file_bytes))
            page_count = len(reader.pages)
        except Exception:
            page_count = max(1, word_count // 400)
    else:
        page_count = max(1, word_count // 400)
    return {"word_count": word_count, "page_count": page_count}
